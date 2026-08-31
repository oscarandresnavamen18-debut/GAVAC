import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def generar_plan_docx():
    doc = Document()

    # Configuración de márgenes estándar (1 pulgada)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilo base
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(51, 51, 51)

    def set_cell_shading(cell, color_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_table_borders(table, color="D3D3D3"):
        tblPr = table._tbl.tblPr
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr.append(borders)

    # Encabezado del Documento
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Plan de Pruebas Manuales")
    r_title.font.size = Pt(16)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(15, 23, 42)

    # Metadatos del Proyecto corregidos con párrafos
    p_meta = doc.add_paragraph()
    p_meta.add_run("• Proyecto: ").bold = True
    p_meta.add_run("GAVAC / Portal Web\n")
    p_meta.add_run("• Versión: ").bold = True
    p_meta.add_run("v1.0.0\n")
    p_meta.add_run("• Fecha: ").bold = True
    p_meta.add_run("30/08/2026\n")
    p_meta.add_run("• Autor: ").bold = True
    p_meta.add_run("QA Lead / Tester\n")

    # Introducción
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("1. Introducción")
    r_h1.font.bold = True
    r_h1.font.size = Pt(13)
    doc.add_paragraph("Este documento describe de forma general la estrategia y los casos de prueba manuales diseñados para validar el correcto funcionamiento de los módulos de autenticación y flujos principales de la aplicación.")

    # Información General y Contexto
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("2. Información General y Contexto")
    r_h2.font.bold = True
    r_h2.font.size = Pt(13)
    
    doc.add_paragraph("• Propósito: Validar el ciclo completo de la aplicación web, desde el inicio de sesión de los usuarios hasta la gestión de transacciones e inventarios.")
    doc.add_paragraph("• Alcance: Login, Registro, Módulo de Ganado, Carrito de Compras. Quedan fuera las librerías externas de terceros.")
    doc.add_paragraph("• Entornos de Prueba:\n- Local: http://localhost:3000 (Desarrollo y pruebas unitarias)\n- Producción: https://gavac-app.vercel.app (Validación de despliegue)")

    # Casos de Prueba (Tabla)
    h3 = doc.add_paragraph()
    r_h3 = h3.add_run("3. Casos de Prueba — Módulo de Autenticación")
    r_h3.font.bold = True
    r_h3.font.size = Pt(13)

    table = doc.add_table(rows=4, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    headers = ["ID", "Título", "Precondiciones", "Pasos", "Resultado Esperado", "Estado"]
    col_widths = [Inches(0.9), Inches(1.1), Inches(1.1), Inches(1.4), Inches(1.2), Inches(0.8)]

    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        hdr_cells[i].width = col_widths[i]
        set_cell_shading(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=80, right=80)
        p = hdr_cells[i].paragraphs[0]
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(255, 255, 255)

    test_data = [
        ("CP-AUTH-001", "Login exitoso con credenciales válidas", "Usuario cliente1@test.com existe en BD de QA.", "1. Navegar a /login.\n2. Ingresar credenciales.\n3. Clic en Iniciar Sesión.", "El sistema redirige al Dashboard (/dashboard).", "Fail"),
        ("CP-AUTH-002", "Error por contraseña incorrecta", "Usuario cliente1@test.com en QA.", "1. Navegar a /login.\n2. Contraseña errónea.\n3. Clic en Ingresar.", "Aparece mensaje en rojo: 'Credenciales inválidas'.", "Pass"),
        ("CP-AUTH-003", "Recuperación de contraseña", "Ninguna.", "1. Ir a /login.\n2. Clic en ¿Olvidaste tu contraseña?.gear.", "Se muestra mensaje de confirmación y envío de enlace.", "Pass")
    ]

    for row_idx, data in enumerate(test_data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_shading(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=80, right=80)
            p = row_cells[col_idx].paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(9)
                if col_idx == 0:
                    run.font.bold = True
                if "Fail" in text:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(185, 28, 28)
                elif text == "Pass":
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(22, 101, 52)

    # Reporte de Defectos (Bugs)
    h4 = doc.add_paragraph()
    r_h4 = h4.add_run("4. Reporte de Defectos (Bugs Encontrados)")
    r_h4.font.bold = True
    r_h4.font.size = Pt(13)

    bug_table = doc.add_table(rows=6, cols=2)
    bug_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(bug_table)

    bug_info = [
        ("Bug ID", "BUG-001"),
        ("Título Descriptivo", "Fallo HTTP 500 y error de JSON al iniciar sesión en Vercel"),
        ("Caso Asociado", "CP-AUTH-001"),
        ("Severidad", "Crítica (Bloqueante)"),
        ("Entorno", "Producción - Google Chrome en Windows 11"),
        ("Descripción y STR", "Al intentar iniciar sesión, la API responde con código 500 indicando revisar variables de entorno y logs. Pasos: 1. Ingresar usuario. 2. Clic en Ingresar.")
    ]

    for idx, (label, val) in enumerate(bug_info):
        row_cells = bug_table.rows[idx].cells
        row_cells[0].text = label
        row_cells[1].text = val
        row_cells[0].width = Inches(2.0)
        row_cells[1].width = Inches(4.5)
        set_cell_shading(row_cells[0], "E2E8F0")
        set_cell_shading(row_cells[1], "F8FAFC")
        set_cell_margins(row_cells[0], top=80, bottom=80, left=100, right=100)
        set_cell_margins(row_cells[1], top=80, bottom=80, left=100, right=100)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.save("Plan_de_Pruebas_Automatico.docx")
    print("[OK] ¡Archivo Word generado con éxito como 'Plan_de_Pruebas_Automatico.docx'!")

if __name__ == "__main__":
    generar_plan_docx()